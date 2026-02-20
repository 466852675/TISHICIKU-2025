import os
import sys


def read_docx(file_path):
    """使用python-docx读取Word文件"""
    if not os.path.exists(file_path):
        print(f"错误: 文件不存在: {file_path}")
        return

    try:
        from docx import Document
    except ImportError:
        print("错误: 未找到python-docx库，请先安装")
        print("安装命令: pip install python-docx")
        return

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"错误: 无法打开Word文档: {file_path}")
        print(f"详细错误: {e}")
        return

    print(f"=== Word文档: {file_path} ===\n")

    try:
        print("=== 文档内容 ===")
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                print(para.text)

        if doc.tables:
            print("\n=== 表格内容 ===")
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\n--- 表格 {table_idx} ---")
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    print(" | ".join(row_data))

    except Exception as e:
        print(f"错误: 读取文档内容时发生错误")
        print(f"详细错误: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python read_docx.py <文件路径>")
        sys.exit(1)

    read_docx(sys.argv[1])
